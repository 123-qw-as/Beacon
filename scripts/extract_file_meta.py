#!/usr/bin/env python
"""提取上传文件的摘要 JSON，供前端展示和 analyst prompt 注入。

用法：python scripts/extract_file_meta.py <file_path>
输出：stdout 一行 JSON
"""
from __future__ import annotations
import json
import sys
from pathlib import Path


def _meta_xlsx(path: Path) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sheets = []
    for i, name in enumerate(wb.sheetnames):
        if i >= 5:
            sheets.append({"name": name, "rows": 0, "cols": 0, "columns": [], "preview": []})
            continue
        ws = wb[name]
        rows = list(ws.iter_rows(max_row=6, values_only=True))
        if not rows:
            sheets.append({"name": name, "rows": 0, "cols": 0, "columns": [], "preview": []})
            continue
        columns = [str(c) if c is not None else "" for c in rows[0]]
        preview = [[str(c) if c is not None else "" for c in row] for row in rows[1:6]]
        # ponytail: max_row in read_only 模式不可靠，用 iter_rows 扫一遍计数
        total_rows = sum(1 for _ in ws.iter_rows(values_only=True))
        sheets.append({
            "name": name, "rows": total_rows, "cols": len(columns),
            "columns": columns, "preview": preview,
        })
    wb.close()
    return {"sheets": sheets}


def _meta_csv(path: Path) -> dict:
    import csv
    with open(path, encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        rows = list(reader)
    if not rows:
        return {"sheets": [{"name": path.name, "rows": 0, "cols": 0, "columns": [], "preview": []}]}
    columns = [str(c) for c in rows[0]]
    preview = [[str(c) for c in row] for row in rows[1:6]]
    return {"sheets": [{"name": path.name, "rows": len(rows), "cols": len(columns),
                         "columns": columns, "preview": preview}]}


def _normalize_math_text(text: str) -> str:
    """规范化 LaTeX PDF 提取的数学文本。

    - LaTeX \\neq 输出为组合短斜线 U+0338 + =，合并为 ≠ (U+2260)
    - LaTeX \\mu 输出为 micro sign U+00B5，统一为希腊 μ (U+03BC)
    """
    text = text.replace("̸=", "≠")     # \neq
    text = text.replace("̸<", "≮")     # \nless
    text = text.replace("̸>", "≯")     # \ngtr
    text = text.replace("̸≤", "≰")     # \nleq
    text = text.replace("̸≥", "≱")     # \ngeq
    text = text.replace("̸∼", "≄")     # \nsim
    text = text.replace("̸≈", "≇")     # \napprox
    text = text.replace("\u00b5", "\u03bc")  # micro sign -> greek mu
    return text


# SymbolMT 等 PUA 字体的字符映射（Word/WPS 生成的 PDF 常见）
_SYMBOL_FONT_MAPS = {
    "SymbolMT": {
        "\uf03d": "=", "\uf02d": "-", "\uf02b": "+", "\uf02f": "/",
        "\uf020": " ", "\uf0d7": "\u00d7", "\uf0b4": "\u2032",
        "\uf0b0": "\u00b0", "\uf070": "\u03c0", "\uf061": "\u03b1",
        "\uf062": "\u03b2", "\uf067": "\u03b3", "\uf071": "\u03b8",
        "\uf06c": "\u03bb", "\uf06d": "\u03bc", "\uf073": "\u03c3",
        "\uf045": "\u2208", "\uf05e": "\u222b", "\uf0d5": "\u221e",
        "\uf0d6": "\u2211",
    },
}

# 标准 Symbol 字体的 PUA 映射（LaTeX/其他工具）
_SYMBOL_FONT_MAPS["Symbol"] = _SYMBOL_FONT_MAPS["SymbolMT"]


def _map_symbol_font(text: str, font: str) -> str:
    """映射 SymbolMT 等 PUA 字体的字符到标准 Unicode。"""
    for font_prefix, char_map in _SYMBOL_FONT_MAPS.items():
        if font.startswith(font_prefix):
            return "".join(char_map.get(c, c) for c in text)
    return text


def _extract_page_with_superscripts(page) -> str:
    """提取页面文本，处理上标和公式排序。

    - 按 x 坐标排序 span（公式中 span 顺序可能乱）
    - SymbolMT 等 PUA 字体字符映射到标准 Unicode
    - 字号明显小于基线的 span 包裹为 ^{...}（上标还原）
    """
    blocks = page.get_text("dict")["blocks"]
    lines_out = []
    for block in blocks:
        # 收集 block 内所有 span
        all_spans = []
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                t = span["text"]
                if not t.strip():
                    continue
                font = span.get("font", "")
                size = round(span["size"], 1)
                x0 = span["bbox"][0]
                t = _map_symbol_font(t, font)
                all_spans.append({"text": t, "font": font, "size": size, "x0": x0})
        if not all_spans:
            continue

        # 判断是否为公式 block：含 PUA/特殊字体的 span 占比较高
        formula_spans = [s for s in all_spans if "Symbol" in s["font"] or "Italic" in s["font"]]
        is_formula = len(formula_spans) > 0 and len(all_spans) > 3

        if is_formula:
            # 公式 block：按 x 排序，baseline 排除正文字体（如 SimSun）
            formula_only = [s for s in all_spans if "SimSun" not in s["font"] and "SimHei" not in s["font"]]
            if formula_only:
                baseline_size = max(set(s["size"] for s in formula_only), key=lambda sz: sum(1 for s in formula_only if s["size"] == sz))
            else:
                baseline_size = max(set(s["size"] for s in all_spans), key=lambda sz: sum(1 for s in all_spans if s["size"] == sz))
            all_spans.sort(key=lambda s: s["x0"])
            parts = []
            for s in all_spans:
                if s["size"] < baseline_size - 1.5:
                    parts.append(f"^{{{s['text']}}}")
                else:
                    parts.append(s["text"])
            lines_out.append("".join(parts))
        else:
            # 正文 block：保持原始行序
            for line in block.get("lines", []):
                spans = [s for s in line.get("spans", []) if s["text"].strip()]
                if not spans:
                    continue
                sizes = [round(s["size"], 1) for s in spans]
                baseline_size = max(set(sizes), key=sizes.count)
                parts = []
                for span in spans:
                    text = _map_symbol_font(span["text"], span.get("font", ""))
                    size = round(span["size"], 1)
                    if size < baseline_size - 1.5:
                        parts.append(f"^{{{text}}}")
                    else:
                        parts.append(text)
                lines_out.append("".join(parts))
    return "\n".join(lines_out)


def _meta_pdf(path: Path) -> dict:
    # 优先使用 PyMuPDF (fitz)，其对数学符号/CID 字体的解码远优于 pypdf
    try:
        import fitz  # pymupdf
        doc = fitz.open(str(path))
        raw = "\n\n".join(_extract_page_with_superscripts(page) for page in doc)
        total_pages = len(doc)
        doc.close()
    except ImportError:
        # 降级到 pypdf（数学符号可能显示为方框，上标会丢失）
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        raw = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        total_pages = len(reader.pages)
    # lone-surrogate 清洗 + 数学符号规范化
    text = raw.encode("utf-8", errors="ignore").decode("utf-8")
    text = _normalize_math_text(text)
    return {"text_excerpt": text[:5000], "total_pages": total_pages}


def _meta_docx(path: Path) -> dict:
    from docx import Document
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return {"text_excerpt": text[:3000], "paragraphs": len(paragraphs), "tables": len(doc.tables)}


def _meta_text(path: Path) -> dict:
    text = path.read_text(encoding="utf-8", errors="ignore")
    return {"text_excerpt": text[:3000], "lines": text.count("\n") + 1}


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "usage: extract_file_meta.py <file_path>"}))
        sys.exit(1)
    path = Path(sys.argv[1])
    if not path.is_file():
        print(json.dumps({"error": f"file not found: {path}"}))
        sys.exit(1)

    suffix = path.suffix.lower()
    type_map = {
        ".xlsx": ("xlsx", _meta_xlsx),
        ".xls": ("xlsx", _meta_xlsx),
        ".csv": ("csv", _meta_csv),
        ".pdf": ("pdf", _meta_pdf),
        ".docx": ("docx", _meta_docx),
        ".txt": ("txt", _meta_text),
        ".md": ("txt", _meta_text),
    }
    if suffix not in type_map:
        print(json.dumps({"error": f"unsupported file type: {suffix}"}))
        sys.exit(1)

    file_type, extractor = type_map[suffix]
    try:
        summary = extractor(path)
    except Exception as e:
        print(json.dumps({"error": f"extraction failed: {e}"}))
        sys.exit(1)

    print(json.dumps({
        "file_type": file_type,
        "filename": path.name,
        "summary": summary,
    }, ensure_ascii=False))


if __name__ == "__main__":
    # UTF-8 输出，避免 Windows GBK 控制台截断中文
    for s in (sys.stdout, sys.stderr):
        if hasattr(s, "reconfigure"):
            s.reconfigure(encoding="utf-8", errors="replace")
    main()
